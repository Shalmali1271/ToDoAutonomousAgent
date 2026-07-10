import os
from datetime import datetime

from docx import Document
from docx.shared import Pt


OUTPUT_DIR = "output"


def create_word_document(title: str, sections: list) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    document = Document()
    # Title section
    title_style = document.add_heading(level=0)
    title_run = title_style.add_run(title)
    title_run.font.size = Pt(22)

    # Metadata section
    document.add_paragraph(
        f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}"
    )

    document.add_paragraph()
    # Sections

    for section in sections:

        heading = section.get("heading", "Section")
        content = section.get("content", "")

        document.add_heading(heading, level=1)

        paragraph = document.add_paragraph()
        paragraph.style.font.size = Pt(12)
        paragraph.add_run(content)

    # Footer section
    document.add_page_break()

    document.add_heading("End of Document", level=2)

    document.add_paragraph(
        "This document was generated automatically by the Autonomous AI Agent."
    )

    # Save section
    filepath = document.save(os.path.join(OUTPUT_DIR, f"{title.lower()}.docx"))

    return filepath